from docx import Document
from docx.shared import Inches
import re
import os

# 读取Markdown文件
with open(r'e:\jianglei\trae\new_payroll\用户手册.md', 'r', encoding='utf-8') as f:
    content = f.read()

# 创建Word文档
doc = Document()

# 解析Markdown内容
lines = content.split('\n')

for line in lines:
    line = line.strip()
    
    # 处理标题
    if line.startswith('#'):
        # 计算标题级别
        level = len(re.match(r'^#+', line).group())
        title_text = line.lstrip('# ')
        
        if level == 1:
            doc.add_heading(title_text, 0)
        elif level == 2:
            doc.add_heading(title_text, 1)
        elif level == 3:
            doc.add_heading(title_text, 2)
        elif level == 4:
            doc.add_heading(title_text, 3)
    
    # 处理图片
    elif line.startswith('!['):
        # 提取图片信息
        img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if img_match:
            img_alt = img_match.group(1)
            img_path = img_match.group(2)
            
            # 检查图片文件是否存在
            full_img_path = os.path.join(r'e:\jianglei\trae\new_payroll', img_path)
            if os.path.exists(full_img_path):
                # 添加图片到文档
                doc.add_picture(full_img_path, width=Inches(6))
                # 添加图片说明
                doc.add_paragraph(img_alt, style='Caption')
            else:
                # 如果图片不存在，添加提示信息
                doc.add_paragraph(f'[图片不存在: {img_path}]')
    
    # 处理列表项
    elif line.startswith('- '):
        # 检查前一个段落是否是列表
        if len(doc.paragraphs) > 0 and doc.paragraphs[-1].style.name == 'List Bullet':
            # 继续列表
            p = doc.paragraphs[-1].insert_paragraph_before('', style='List Bullet')
            p.add_run(line.lstrip('- '))
        else:
            # 开始新列表
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line.lstrip('- '))
    
    # 处理数字列表项
    elif re.match(r'^\d+\. ', line):
        # 检查前一个段落是否是数字列表
        if len(doc.paragraphs) > 0 and doc.paragraphs[-1].style.name == 'List Number':
            # 继续列表
            p = doc.paragraphs[-1].insert_paragraph_before('', style='List Number')
            p.add_run(re.sub(r'^\d+\. ', '', line))
        else:
            # 开始新数字列表
            p = doc.add_paragraph(style='List Number')
            p.add_run(re.sub(r'^\d+\. ', '', line))
    
    # 处理普通段落
    elif line:
        doc.add_paragraph(line)
    
    # 处理空行
    else:
        # 添加空段落
        doc.add_paragraph()

# 保存Word文档
doc.save('用户手册.docx')
print('Markdown已成功转换为Word文档！')